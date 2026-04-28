"""
ファイル変換エンジン。
多様なファイル形式を Markdown テキストに変換する。
"""

import csv
import io
import json
import logging
import uuid
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path

from app.services.text_normalizer import normalize_text

logger = logging.getLogger(__name__)


class FileType(str, Enum):
    MD = "md"
    TXT = "txt"
    CSV = "csv"
    JSON = "json"
    PDF = "pdf"
    PPTX = "pptx"
    XLSX = "xlsx"
    DOCX = "docx"
    HTML = "html"
    PNG = "png"
    JPEG = "jpeg"
    JPG = "jpg"


@dataclass
class ExtractedImage:
    uuid: str
    data: bytes
    mime_type: str
    caption: str | None = None
    page: int | None = None


@dataclass
class ConversionResult:
    markdown: str
    images: list[ExtractedImage] = field(default_factory=list)
    metadata_hints: dict[str, str | int | None] = field(default_factory=dict)


def detect_file_type(filename: str) -> FileType:
    """ファイル名から FileType を判定する。"""
    ext = Path(filename).suffix.lstrip(".").lower()
    try:
        return FileType(ext)
    except ValueError as e:
        raise ValueError(f"Unsupported file type: {ext}") from e


async def convert_file(file_path: str, filename: str) -> ConversionResult:
    """
    ファイルを Markdown に変換する。
    ファイルタイプに応じて適切な変換関数を呼び出す。
    """
    import asyncio

    file_type = detect_file_type(filename)

    converters = {
        FileType.MD: _convert_markdown,
        FileType.TXT: _convert_text,
        FileType.CSV: _convert_csv,
        FileType.JSON: _convert_json,
        FileType.PDF: _convert_pdf,
        FileType.PPTX: _convert_pptx,
        FileType.XLSX: _convert_xlsx,
        FileType.DOCX: _convert_docx,
        FileType.HTML: _convert_html,
        FileType.PNG: _convert_image,
        FileType.JPEG: _convert_image,
        FileType.JPG: _convert_image,
    }

    converter_fn = converters.get(file_type)
    if converter_fn is None:
        raise ValueError(f"No converter for file type: {file_type}")

    result = await asyncio.to_thread(converter_fn, file_path, filename)
    # Apply NFKC normalization to all converted text
    result.markdown = normalize_text(result.markdown)
    return result


def _convert_markdown(file_path: str, filename: str) -> ConversionResult:
    """MD ファイル: パススルー"""
    text = Path(file_path).read_text(encoding="utf-8")
    return ConversionResult(markdown=text, metadata_hints={"title": filename})


def _convert_text(file_path: str, filename: str) -> ConversionResult:
    """TXT ファイル: パススルー"""
    text = Path(file_path).read_text(encoding="utf-8")
    return ConversionResult(markdown=text, metadata_hints={"title": filename})


def _convert_csv(file_path: str, filename: str) -> ConversionResult:
    """CSV ファイル: Markdown テーブルに変換"""
    text = Path(file_path).read_text(encoding="utf-8")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)

    if not rows:
        return ConversionResult(
            markdown="(empty CSV)", metadata_hints={"title": filename}
        )

    # Build markdown table
    headers = rows[0]
    md_lines = [
        f"# {filename}",
        "",
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows[1:]:
        # Pad row if needed
        padded = row + [""] * (len(headers) - len(row))
        md_lines.append("| " + " | ".join(padded[: len(headers)]) + " |")

    return ConversionResult(
        markdown="\n".join(md_lines),
        metadata_hints={"title": filename, "row_count": len(rows) - 1},
    )


def _convert_json(file_path: str, filename: str) -> ConversionResult:
    """JSON ファイル: 構造解析して Markdown に変換"""
    text = Path(file_path).read_text(encoding="utf-8")
    data = json.loads(text)

    if isinstance(data, list) and data and isinstance(data[0], dict):
        # Array of objects -> table
        headers = list(data[0].keys())
        md_lines = [
            f"# {filename}",
            "",
            "| " + " | ".join(headers) + " |",
            "| " + " | ".join(["---"] * len(headers)) + " |",
        ]
        for item in data:
            row = [str(item.get(h, "")) for h in headers]
            md_lines.append("| " + " | ".join(row) + " |")
        return ConversionResult(
            markdown="\n".join(md_lines),
            metadata_hints={"title": filename, "row_count": len(data)},
        )
    else:
        # Generic JSON -> code block
        formatted = json.dumps(data, ensure_ascii=False, indent=2)
        md = f"# {filename}\n\n```json\n{formatted}\n```"
        return ConversionResult(markdown=md, metadata_hints={"title": filename})


def _convert_pdf(file_path: str, filename: str) -> ConversionResult:
    """PDF ファイル: pymupdf4llm で変換"""
    try:
        import pymupdf4llm

        markdown = pymupdf4llm.to_markdown(file_path)
        return ConversionResult(
            markdown=f"# {filename}\n\n{markdown}",
            metadata_hints={"title": filename},
        )
    except ImportError:
        logger.warning("pymupdf4llm not installed, using fallback")
        return ConversionResult(
            markdown=f"# {filename}\n\n[変換失敗: pymupdf4llm がインストールされていません]",
            metadata_hints={"title": filename},
        )
    except Exception as e:
        logger.error("PDF conversion failed: %s", e)
        return ConversionResult(
            markdown=f"# {filename}\n\n[変換失敗: {e}]",
            metadata_hints={"title": filename},
        )


def _convert_pptx(file_path: str, filename: str) -> ConversionResult:
    """PPTX ファイル: pptx2md で変換"""
    try:
        import tempfile
        from pathlib import Path
        from pptx2md import ConversionConfig, convert as pptx_convert

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "output.md"
            image_dir = Path(tmp_dir) / "images"
            config = ConversionConfig(
                pptx_path=Path(file_path),
                output_path=output_path,
                image_dir=image_dir,
                disable_image=True,
            )
            pptx_convert(config)
            markdown = (
                output_path.read_text(encoding="utf-8") if output_path.exists() else ""
            )
        return ConversionResult(
            markdown=f"# {filename}\n\n{markdown}",
            metadata_hints={"title": filename},
        )
    except ImportError:
        logger.warning("pptx2md not installed, using fallback")
        return ConversionResult(
            markdown=f"# {filename}\n\n[変換失敗: pptx2md がインストールされていません]",
            metadata_hints={"title": filename},
        )
    except Exception as e:
        logger.error("PPTX conversion failed: %s", e)
        return ConversionResult(
            markdown=f"# {filename}\n\n[変換失敗: {e}]",
            metadata_hints={"title": filename},
        )


def _convert_xlsx(file_path: str, filename: str) -> ConversionResult:
    """XLSX ファイル: openpyxl でシートごとにMarkdownテーブルへ変換"""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        md_parts: list[str] = [f"# {filename}"]

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            md_parts.append(f"\n## {sheet_name}\n")

            # Build markdown table
            header = rows[0]
            cols = [str(c) if c is not None else "" for c in header]
            md_parts.append("| " + " | ".join(cols) + " |")
            md_parts.append("| " + " | ".join("---" for _ in cols) + " |")

            for row in rows[1:]:
                cells = [str(c) if c is not None else "" for c in row]
                # Pad or trim to match header column count
                while len(cells) < len(cols):
                    cells.append("")
                md_parts.append("| " + " | ".join(cells[: len(cols)]) + " |")

        wb.close()

        markdown = "\n".join(md_parts)
        return ConversionResult(
            markdown=markdown,
            metadata_hints={"title": filename},
        )
    except Exception as e:
        logger.error("XLSX conversion failed: %s", e)
        return ConversionResult(
            markdown=f"# {filename}\n\n[変換失敗: {e}]",
            metadata_hints={"title": filename},
        )


def _convert_docx(file_path: str, filename: str) -> ConversionResult:
    """DOCX ファイル: python-docx で変換"""
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        images: list[ExtractedImage] = []
        md_lines = [f"# {filename}", ""]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                md_lines.append("")
                continue

            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name:
                md_lines.append(f"## {text}")
            elif "heading 2" in style_name:
                md_lines.append(f"### {text}")
            elif "heading 3" in style_name:
                md_lines.append(f"#### {text}")
            elif "list" in style_name:
                md_lines.append(f"- {text}")
            else:
                md_lines.append(text)

        # Extract tables
        for table in doc.tables:
            md_lines.append("")
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    md_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

        return ConversionResult(
            markdown="\n".join(md_lines),
            images=images,
            metadata_hints={"title": filename},
        )
    except ImportError:
        return ConversionResult(
            markdown=f"# {filename}\n\n[変換失敗: python-docx がインストールされていません]",
            metadata_hints={"title": filename},
        )
    except Exception as e:
        logger.error("DOCX conversion failed: %s", e)
        return ConversionResult(
            markdown=f"# {filename}\n\n[変換失敗: {e}]",
            metadata_hints={"title": filename},
        )


def _convert_html(file_path: str, filename: str) -> ConversionResult:
    """HTML ファイル: beautifulsoup4 + markdownify で変換"""
    try:
        from bs4 import BeautifulSoup
        from markdownify import markdownify as md_convert

        html = Path(file_path).read_text(encoding="utf-8")
        soup = BeautifulSoup(html, "html.parser")

        # Extract title
        title = filename
        title_tag = soup.find("title")
        if title_tag and title_tag.string:
            title = title_tag.string.strip()

        # Remove script and style tags
        for tag in soup.find_all(["script", "style"]):
            tag.decompose()

        markdown = md_convert(str(soup), heading_style="ATX")
        return ConversionResult(
            markdown=markdown,
            metadata_hints={"title": title},
        )
    except ImportError:
        return ConversionResult(
            markdown=f"# {filename}\n\n[変換失敗: beautifulsoup4/markdownify がインストールされていません]",
            metadata_hints={"title": filename},
        )
    except Exception as e:
        logger.error("HTML conversion failed: %s", e)
        return ConversionResult(
            markdown=f"# {filename}\n\n[変換失敗: {e}]",
            metadata_hints={"title": filename},
        )


def _convert_image(file_path: str, filename: str) -> ConversionResult:
    """PNG/JPEG: 画像マーカーを埋め込み、Vision 解析は後続で実行"""
    image_uuid = str(uuid.uuid4())
    image_data = Path(file_path).read_bytes()

    ext = Path(filename).suffix.lower()
    mime_map = {".png": "image/png", ".jpeg": "image/jpeg", ".jpg": "image/jpeg"}
    mime_type = mime_map.get(ext, "image/png")

    extracted = ExtractedImage(
        uuid=image_uuid,
        data=image_data,
        mime_type=mime_type,
        caption=filename,
    )

    markdown = f"# {filename}\n\n[[IMAGE:{image_uuid}]]\n\n*画像ファイル: {filename}*"

    return ConversionResult(
        markdown=markdown,
        images=[extracted],
        metadata_hints={"title": filename},
    )
